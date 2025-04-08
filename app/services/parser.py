import io
from typing import List
from fastapi import UploadFile, HTTPException, status
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import logging
import os

# 导入特定的加载器 - 确保这些与 Langchain 0.3 兼容
# 查阅 Langchain 0.3 文档以获取正确的加载器类和导入路径。
# 这些可能位于 langchain.document_loaders 或类似位置。
from langchain_community.document_loaders import PyPDFLoader, UnstructuredMarkdownLoader, TextLoader
# 对于 DOCX, UnstructuredFileLoader 比较常用, 或者可能需要直接使用 python-docx
# from langchain_community.document_loaders import UnstructuredFileLoader # 示例
import docx # 直接使用 python-docx 作为替代方案

# 允许的内容类型及其对应的文件扩展名
ALLOWED_CONTENT_TYPES = {
    "application/pdf": ".pdf",
    "text/plain": ".txt",
    "text/markdown": ".md",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx"
}

# 文本分割参数
CHUNK_SIZE = 1000 # 块大小
CHUNK_OVERLAP = 150 # 块重叠大小

def _load_pdf(file_content: bytes, filename: str) -> List[Document]:
    """加载 PDF 文件内容。"""
    # PyPDFLoader 通常需要文件路径。我们需要将字节保存到临时文件中。
    # 或者, 检查 Langchain 0.3 的 PyPDFLoader 是否能直接处理字节。
    # 如果需要路径, 这是常用的模式:
    tmp_path = None # Initialize tmp_path
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpfile:
            tmpfile.write(file_content)
            tmp_path = tmpfile.name
        loader = PyPDFLoader(tmp_path)
        # Langchain 0.3 的加载器可能直接返回 Document, 或需要调用 .load() 方法。
        # 根据具体版本的 API 进行调整。
        docs = loader.load() # 或者可能是 `loader.load_and_split()` (如果可用)
        for doc in docs:
            doc.metadata["source"] = filename # 将原始文件名添加到元数据
        return docs
    except Exception as e:
        logging.error(f"加载 PDF {filename} 时出错: {e}", exc_info=True)
        # Don't raise HTTPException here, let the caller handle it
        raise IOError(f"解析 PDF 失败: {e}") from e
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError as remove_err:
                 logging.warning(f"清理临时 PDF 文件 {tmp_path} 失败: {remove_err}")

def _load_markdown(file_content: bytes, filename: str) -> List[Document]:
    """加载 Markdown 文件内容。"""
    # 与 PDF 类似, 加载器通常需要文件路径。
    tmp_path = None # Initialize tmp_path
    try:
        content = file_content.decode('utf-8') # 假设 Markdown 文件使用 UTF-8 编码
        # UnstructuredMarkdownLoader 可能接受内容或需要路径
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".md", encoding='utf-8') as tmpfile:
            tmpfile.write(content)
            tmp_path = tmpfile.name
        loader = UnstructuredMarkdownLoader(tmp_path)
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = filename
        return docs
    except Exception as e:
        logging.error(f"加载 Markdown {filename} 时出错: {e}", exc_info=True)
        raise IOError(f"解析 Markdown 失败: {e}") from e
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError as remove_err:
                 logging.warning(f"清理临时 Markdown 文件 {tmp_path} 失败: {remove_err}")

def _load_text(file_content: bytes, filename: str) -> List[Document]:
    """加载纯文本文件内容。"""
    # TextLoader 可能也需要路径, 或者有时可以处理原始文本。
    tmp_path = None # Initialize tmp_path
    try:
        content = file_content.decode('utf-8') # 假设文本文件使用 UTF-8 编码
        # 如果 TextLoader 需要路径:
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding='utf-8') as tmpfile:
            tmpfile.write(content)
            tmp_path = tmpfile.name
        loader = TextLoader(tmp_path, encoding='utf-8')
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = filename
        return docs
        # 如果它可以处理内容 (备选方案):
        # text_splitter = RecursiveCharacterTextSplitter(...) # 定义你的分割器
        # texts = text_splitter.split_text(content)
        # return [Document(page_content=t, metadata={"source": filename}) for t in texts]
    except Exception as e:
        logging.error(f"加载文本文件 {filename} 时出错: {e}", exc_info=True)
        raise IOError(f"解析文本文件失败: {e}") from e
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError as remove_err:
                 logging.warning(f"清理临时文本文件 {tmp_path} 失败: {remove_err}")

def _load_docx(file_path: str, filename: str) -> List[Document]:
    """加载 DOCX 文件内容 (从路径加载)。"""
    logger = logging.getLogger(__name__) # Get logger instance
    try:
        # Log file existence before attempting to load
        file_exists = os.path.exists(file_path)
        logger.debug(f"尝试加载 DOCX: '{filename}'. 路径: '{file_path}'. 文件是否存在: {file_exists}")
        if not file_exists:
             raise IOError(f"文件在尝试加载时未找到: {file_path}")
             
        # Load directly from the file path
        document = docx.Document(file_path)
        full_text = "\n".join([para.text for para in document.paragraphs if para.text])
        if not full_text:
            logger.warning(f"DOCX 文件 {filename} 解析后内容为空。")
            return [] # Return empty list if no text found
        return [Document(page_content=full_text, metadata={"source": filename})]
    except Exception as e:
        logger.error(f"加载 DOCX {filename} (路径: {file_path}) 时出错: {e}", exc_info=True)
        # Ensure specific error type is raised for task handling
        if isinstance(e, (IOError, FileNotFoundError)):
             raise IOError(f"解析 DOCX IO 错误: {e}") from e
        else:
             # Wrap other errors (like BadZipFile if it happens again)
             raise IOError(f"解析 DOCX 文件失败: {e}") from e

def parse_file_from_path_and_split(file_path: str, original_filename: str) -> List[Document]:
    """从文件路径同步加载、解析和分割文档。"""
    logger = logging.getLogger(__name__) # Get logger instance
    logger.debug(f"开始同步解析文件: {original_filename} (路径: {file_path})")
    _, ext = os.path.splitext(original_filename)
    ext = ext.lower()
    
    content_type_map = {
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    
    content_type = content_type_map.get(ext)
    
    if not content_type:
        raise ValueError(f"根据扩展名无法确定支持的文件类型: {ext}")
        
    try:
        with open(file_path, "rb") as f:
            file_content = f.read()
    except OSError as e:
        logger.error(f"读取文件 {file_path} 时出错: {e}")
        raise IOError(f"无法读取文件: {e}") from e
        
    if not file_content:
        logger.warning(f"文件 {file_path} 为空。")
        return []

    raw_docs: List[Document] = []
    try:
        if content_type == "application/pdf":
            raw_docs = _load_pdf(file_content, original_filename)
        elif content_type == "text/markdown":
            raw_docs = _load_markdown(file_content, original_filename)
        elif content_type == "text/plain":
            raw_docs = _load_text(file_content, original_filename)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raw_docs = _load_docx(file_path, original_filename)
        else:
             # Should not happen due to check above
             raise ValueError(f"内部错误：未处理的内容类型 {content_type}")
             
    except IOError as e: # Catch specific IOErrors from loaders
        logger.error(f"加载/解析 {original_filename} 时出错: {e}")
        # Re-raise or handle as needed by the task
        raise
    except Exception as e: # Catch other unexpected errors during loading
        logger.exception(f"加载/解析 {original_filename} 时发生意外错误: {e}")
        raise RuntimeError(f"解析时发生意外错误: {e}") from e

    if not raw_docs:
         logger.warning(f"文档 {original_filename} 已加载但未产生任何内容。")
         return [] # Return empty list, not an error

    # Split documents
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        add_start_index=True,
    )

    all_splits = []
    for doc in raw_docs:
        if isinstance(doc.page_content, str):
            try:
                 splits = text_splitter.split_documents([doc])
                 # Ensure metadata is copied and source is present
                 for split in splits:
                     split.metadata = {**split.metadata, 'source': original_filename}
                     split.metadata.pop('start_index', None)
                 all_splits.extend(splits)
            except Exception as split_e:
                 logger.error(f"分割 {original_filename} 的块时出错: {split_e}", exc_info=True)
                 # Optionally skip this doc or handle error
        else:
            logger.warning(f"在 {original_filename} 中跳过非字符串内容的文档块")
            continue

    if not all_splits:
        logger.warning(f"文档 {original_filename} 内容无法分割成可处理的块。")
        return [] # Return empty list

    logger.info(f"成功将 '{original_filename}' (来自路径 {file_path}) 解析并分割成 {len(all_splits)} 个块。")
    return all_splits

async def parse_uploaded_file_and_split(file: UploadFile) -> List[Document]:
    """根据上传文件的内容类型解析文件并将其分割成块 (供API使用)。"""
    # ... (这个函数的实现保持不变，它调用同步的 _load_... 函数)
    # ... 只是重命名以明确其用途
    content_type = file.content_type
    filename = file.filename

    if content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件类型: {content_type}。允许的类型: {list(ALLOWED_CONTENT_TYPES.keys())}"
        )

    file_content = await file.read() # Read content once for loaders that need it
    if not file_content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="文件为空")

    # Need a temporary file for _load_docx if it now requires a path
    # Or revert _load_docx to use stream for the async path?
    # Let's keep _load_docx taking path, and create temp file here for it.
    
    temp_file_for_docx = None
    raw_docs: List[Document] = []
    try:
        if content_type == "application/pdf":
            raw_docs = _load_pdf(file_content, filename)
        elif content_type == "text/markdown":
            raw_docs = _load_markdown(file_content, filename)
        elif content_type == "text/plain":
            raw_docs = _load_text(file_content, filename)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
             # Create a temporary file to pass its path to _load_docx
             import tempfile
             with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
                  tmpfile.write(file_content)
                  temp_file_for_docx = tmpfile.name
             raw_docs = _load_docx(temp_file_for_docx, filename)
        else:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="未处理的内容类型")

    except IOError as e: # Catch specific IOErrors from loaders
         raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"解析文件失败: {e}")
    except Exception as e: # Catch other unexpected errors
        logger.exception(f"解析上传文件 {filename} 时发生意外错误: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"处理文件时发生意外错误。")
    finally:
         # Clean up the temp file created for docx in the async path
         if temp_file_for_docx and os.path.exists(temp_file_for_docx):
              try:
                   os.remove(temp_file_for_docx)
              except OSError as remove_err:
                   logger.warning(f"清理上传路径的临时 DOCX 文件 {temp_file_for_docx} 失败: {remove_err}")

    if not raw_docs:
         raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="文档已加载但未产生任何内容。")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        add_start_index=True,
    )

    all_splits = []
    for doc in raw_docs:
        if isinstance(doc.page_content, str):
            try:
                splits = text_splitter.split_documents([doc])
                for split in splits:
                     split.metadata = {**split.metadata, 'source': filename}
                     split.metadata.pop('start_index', None)
                all_splits.extend(splits)
            except Exception as split_e:
                 logger.error(f"分割上传文件 {filename} 的块时出错: {split_e}", exc_info=True)
        else:
            logger.warning(f"在上传文件 {filename} 中跳过非字符串内容的文档块")
            continue

    if not all_splits:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="文档内容无法分割成可处理的块。")

    logger.info(f"成功将上传文件 '{filename}' 解析并分割成 {len(all_splits)} 个块。")
    return all_splits 